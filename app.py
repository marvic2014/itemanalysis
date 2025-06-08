from flask import Flask, render_template, request, send_file, redirect, url_for, flash
import pandas as pd
import math
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.secret_key = 'supersecret'
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route('/')
def home():
    return render_template("home.html")

@app.route('/upload', methods=['GET', 'POST'])
def upload():
    if request.method == 'POST':
        percent = request.form.get('percentage')
        file = request.files.get('file')

        if not file or percent == "Select %":
            flash("Missing file or percentage selection.")
            return redirect(request.url)

        filepath = os.path.join(UPLOAD_FOLDER, secure_filename(file.filename))
        file.save(filepath)

        result = process_excel_web(filepath, percent)
        return render_template("result.html", result=result)

    return render_template("upload.html")

def process_excel_web(filepath, percent):
    try:
        df = pd.read_excel(filepath)
        num_items = len(df.columns) - df.columns.get_loc("Key Version") - 1
        sample_percent = int(percent.replace('%', ''))

        df_cleaned = df.dropna(how='all')
        df_sorted = df_cleaned.sort_values(by="Num Correct", ascending=False)
        row_count = len(df_sorted)
        sample_size = math.ceil(row_count * (sample_percent / 100.0))

        question_columns = list(df.columns[df.columns.get_loc("Key Version") + 1:])

        top_data = df_sorted.head(sample_size)
        top_data = top_data.drop(columns=['First Name', 'Last Name'], errors='ignore')

        bottom_data = df_sorted.tail(sample_size).iloc[::-1]
        bottom_data = bottom_data.drop(columns=['First Name', 'Last Name'], errors='ignore')

        ug_values = top_data[question_columns].sum()
        lg_values = bottom_data[question_columns].sum()

        difficulty_index = (ug_values + lg_values) / (2 * sample_size)
        discrimination_index = (ug_values - lg_values) / sample_size

        difficulty_remarks = ["Too Easy" if d >= 0.80 else "Too Difficult" if d <= 0.30 else "Acceptable" for d in difficulty_index]
        discrimination_remarks = ["Retained" if d >= 0.40 else "Revised" if d >= 0.20 else "Discard" for d in discrimination_index]

        item_df = pd.DataFrame({
            "Item Number": question_columns,
            "UG": ug_values.values,
            "LG": lg_values.values,
            "Difficulty Index": difficulty_index.values.round(2),
            "Discrimination Index": discrimination_index.values.round(2),
            "Difficulty Remarks": difficulty_remarks,
            "Discrimination Remarks": discrimination_remarks,
        })

        mean_score = df_sorted["Num Correct"].mean()
        std_dev_score = df_sorted["Num Correct"].std()
        mean_percent_score = (mean_score / num_items) * 100

        summary_df = pd.DataFrame({
            "Metric": ["Mean Score", "Standard Deviation", "Mean Percentage Score"],
            "Value": [round(mean_score, 2), round(std_dev_score, 2), f"{round(mean_percent_score, 2)}%"]
        })

        excel_output = os.path.join(UPLOAD_FOLDER, "Item_Analysis_Report.xlsx")
        with pd.ExcelWriter(excel_output, engine='openpyxl') as writer:
            top_data.to_excel(writer, sheet_name="Top Students", index=False)
            bottom_data.to_excel(writer, sheet_name="Bottom Students", index=False)
            item_df.to_excel(writer, sheet_name="Item Analysis", index=False)
            summary_df.to_excel(writer, sheet_name="Summary Statistics", index=False)

        word_output = os.path.join(UPLOAD_FOLDER, "Item_Analysis.docx")
        doc = Document()
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(13)
        header_lines = [
            "Republic of the Philippines",
            "Department of Education",
            "Region _____",
            "Schools Division Office ________________",
            "SCHOOL NAME",
            "SCHOOL ADDRESS"
        ]
        for line in header_lines:
            para = doc.add_paragraph(line)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.runs[0]
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

        doc.add_heading("Item Analysis Report", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("\n\nSubject:\t\t\t\t\t\tQuarter:\n\nPrepared by:\t\t\t\t\t\tChecked by:\n\nApproved by:\t\t\t\t\t\tNoted by:\n\n")

        def add_dataframe_to_doc(df, doc, heading):
            doc.add_paragraph(heading)
            table = doc.add_table(rows=1, cols=len(df.columns))
            try:
                table.style = 'Table Grid'
            except:
                table.style = 'Normal'
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(df.columns):
                hdr_cells[i].text = str(col_name)
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
            doc.add_paragraph("")

        add_dataframe_to_doc(item_df, doc, "Item Analysis Table")
        add_dataframe_to_doc(summary_df, doc, "Summary Statistics")
        doc.save(word_output)

        return {"excel": os.path.basename(excel_output), "word": os.path.basename(word_output)}

    except Exception as e:
        return {"error": str(e)}

@app.route('/download/<filename>')
def download(filename):
    return send_file(os.path.join(UPLOAD_FOLDER, filename), as_attachment=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
